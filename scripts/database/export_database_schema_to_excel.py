#!/usr/bin/env python3
"""
Export PreTech-NIDS Database Schema to Excel and Word
Creates formatted Excel and Word files with all database collections and their schemas
"""

try:
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
    from openpyxl.utils import get_column_letter
    import os
    OPENPYXL_AVAILABLE = True
except ImportError:
    print("❌ openpyxl not available, creating CSV files instead...")
    OPENPYXL_AVAILABLE = False
    import csv

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.shared import OxmlElement, qn
    WORD_AVAILABLE = True
except ImportError:
    print("⚠️ python-docx not available, Word document will not be created")
    WORD_AVAILABLE = False

try:
    import csv
    CSV_AVAILABLE = True
except ImportError:
    print("❌ CSV module not available")
    CSV_AVAILABLE = False

def create_schema_data():
    """Create the complete database schema data"""
    
    schemas = {
        "Users Collection": {
            "Field Name": ["_id", "uid", "display_name", "email", "phone_number", "profile_pic", "role", "created_time", "last_login", "is_active"],
            "Data Type": ["ObjectId", "string", "string", "string", "string", "string", "string", "timestamp", "timestamp", "boolean"],
            "Constraints": ["Primary Key (PK)", "NOT NULL", "NOT NULL", "NOT NULL", "NOT NULL", "-", "NOT NULL", "NOT NULL", "-", "NOT NULL"],
            "Description": [
                "MongoDB auto-generated unique identifier",
                "User unique identifier",
                "User display name",
                "User email address",
                "User phone number",
                "Profile picture URL path",
                "User role (admin, analyst, user)",
                "User account creation timestamp",
                "Last login time",
                "Account activation status"
            ]
        },
        
        "Detection Reports Collection": {
            "Field Name": ["_id", "model", "input", "output", "timestamp", "type", "interface", "result", "src_ip", "dst_ip", "dst_port", "protocol"],
            "Data Type": ["ObjectId", "string", "array[float]", "object", "string", "string", "string", "object", "string", "string", "integer", "string"],
            "Constraints": ["Primary Key (PK)", "NOT NULL", "NOT NULL", "NOT NULL", "NOT NULL", "NOT NULL", "-", "NOT NULL", "-", "-", "-", "-"],
            "Description": [
                "MongoDB auto-generated unique identifier",
                "Machine learning model name used",
                "Input feature vector (77-dimensional features)",
                "Model prediction results and confidence",
                "Detection timestamp (ISO format)",
                "Detection type (manual_testing/real_time_detection)",
                "Network interface name",
                "Prediction result details (prediction, probability)",
                "Source IP address from packet",
                "Destination IP address from packet",
                "Destination port number",
                "Protocol type (TCP/UDP/ICMP)"
            ]
        },
        
        "Alerts Collection": {
            "Field Name": ["_id", "rule_id", "alert_type", "level", "title", "message", "source_ip", "destination_ip", "target_port", "protocol", "model", "confidence", "timestamp", "acknowledged", "acknowledged_by", "resolved", "resolved_by"],
            "Data Type": ["ObjectId", "string", "string", "string", "string", "string", "string", "string", "integer", "string", "string", "float", "string", "boolean", "string", "boolean", "string"],
            "Constraints": ["Primary Key (PK)", "NOT NULL", "NOT NULL", "NOT NULL", "NOT NULL", "NOT NULL", "-", "-", "-", "-", "-", "-", "NOT NULL", "NOT NULL", "-", "NOT NULL", "-"],
            "Description": [
                "MongoDB auto-generated unique identifier",
                "Alert rule identifier",
                "Alert type (threat_detected, anomaly_detected, etc.)",
                "Alert level (critical, high, medium, low, info)",
                "Alert title",
                "Alert detailed information",
                "Source IP address",
                "Destination IP address",
                "Target port number",
                "Protocol type",
                "Model name that triggered the alert",
                "Detection confidence",
                "Alert generation timestamp",
                "Whether acknowledged",
                "Acknowledging user ID",
                "Whether resolved",
                "Resolving user ID"
            ]
        },
        
        "Alert Rules Collection": {
            "Field Name": ["_id", "id", "name", "description", "alert_type", "conditions", "actions", "enabled", "threshold", "time_window", "created_at", "updated_at"],
            "Data Type": ["ObjectId", "string", "string", "string", "string", "object", "array[string]", "boolean", "float", "integer", "string", "string"],
            "Constraints": ["Primary Key (PK)", "NOT NULL", "NOT NULL", "NOT NULL", "NOT NULL", "NOT NULL", "NOT NULL", "NOT NULL", "-", "-", "NOT NULL", "NOT NULL"],
            "Description": [
                "MongoDB auto-generated unique identifier",
                "Rule unique identifier",
                "Rule name",
                "Rule description",
                "Alert type",
                "Trigger condition configuration",
                "Action list to execute",
                "Whether rule is enabled",
                "Trigger threshold",
                "Time window (minutes)",
                "Rule creation time",
                "Rule update time"
            ]
        },
        
        "Alert History Collection": {
            "Field Name": ["_id", "alert_id", "action", "user_id", "timestamp", "notes", "previous_state", "new_state"],
            "Data Type": ["ObjectId", "string", "string", "string", "string", "string", "object", "object"],
            "Constraints": ["Primary Key (PK)", "NOT NULL", "NOT NULL", "NOT NULL", "NOT NULL", "-", "NOT NULL", "NOT NULL"],
            "Description": [
                "MongoDB auto-generated unique identifier",
                "Original alert ID",
                "Action performed (acknowledge, resolve, escalate, etc.)",
                "User ID who performed the action",
                "Action execution time",
                "Notes information",
                "State before operation",
                "State after operation"
            ]
        },
        
        "Attack Locations Collection": {
            "Field Name": ["_id", "timestamp", "source_ip", "location", "attack_details", "country", "country_code", "latitude", "longitude"],
            "Data Type": ["ObjectId", "string", "string", "object", "object", "string", "string", "float", "float"],
            "Constraints": ["Primary Key (PK)", "NOT NULL", "NOT NULL", "NOT NULL", "NOT NULL", "NOT NULL", "NOT NULL", "NOT NULL", "NOT NULL"],
            "Description": [
                "MongoDB auto-generated unique identifier",
                "Attack detection timestamp",
                "Source IP address of the attack",
                "Geographic location information",
                "Attack details including model and prediction",
                "Country name",
                "Country code (ISO format)",
                "Latitude coordinate",
                "Longitude coordinate"
            ]
        },
        
        "Geo Cache Collection": {
            "Field Name": ["_id", "ip", "location_data", "cached_at"],
            "Data Type": ["ObjectId", "string", "object", "string"],
            "Constraints": ["Primary Key (PK)", "NOT NULL", "NOT NULL", "NOT NULL"],
            "Description": [
                "MongoDB auto-generated unique identifier",
                "IP address",
                "Geographic location data",
                "Cache timestamp"
            ]
        },
        
        "PCAP Analyses Collection": {
            "Field Name": ["_id", "filename", "file_hash", "file_size", "total_packets", "packet_analysis", "threat_analysis", "timestamp", "processing_time", "status"],
            "Data Type": ["ObjectId", "string", "string", "integer", "integer", "object", "object", "string", "float", "string"],
            "Constraints": ["Primary Key (PK)", "NOT NULL", "NOT NULL", "NOT NULL", "NOT NULL", "NOT NULL", "NOT NULL", "NOT NULL", "NOT NULL", "NOT NULL"],
            "Description": [
                "MongoDB auto-generated unique identifier",
                "PCAP filename",
                "File MD5 hash value",
                "File size (bytes)",
                "Total packet count",
                "Packet analysis results",
                "Threat analysis and assessment",
                "Analysis completion time",
                "Processing time (seconds)",
                "Analysis status (completed, failed, processing)"
            ]
        },
        
        "PCAP Reports Collection": {
            "Field Name": ["_id", "analysis_id", "report_type", "content", "generated_at", "format"],
            "Data Type": ["ObjectId", "string", "string", "object", "string", "string"],
            "Constraints": ["Primary Key (PK)", "NOT NULL", "NOT NULL", "NOT NULL", "NOT NULL", "NOT NULL"],
            "Description": [
                "MongoDB auto-generated unique identifier",
                "Associated analysis ID",
                "Report type (summary, detailed, threat_analysis)",
                "Report content",
                "Report generation time",
                "Report format (json, pdf, html)"
            ]
        },
        
        "Password Resets Collection": {
            "Field Name": ["_id", "email", "token", "expires_at", "used", "created_at"],
            "Data Type": ["ObjectId", "string", "string", "timestamp", "boolean", "timestamp"],
            "Constraints": ["Primary Key (PK)", "NOT NULL", "NOT NULL", "NOT NULL", "NOT NULL", "NOT NULL"],
            "Description": [
                "MongoDB auto-generated unique identifier",
                "User email address",
                "Reset token",
                "Expiration time",
                "Whether used",
                "Creation time"
            ]
        },
        
        "Registration Verifications Collection": {
            "Field Name": ["_id", "email", "otp_code", "expires_at", "verified", "created_at", "verified_at"],
            "Data Type": ["ObjectId", "string", "string", "timestamp", "boolean", "timestamp", "timestamp"],
            "Constraints": ["Primary Key (PK)", "NOT NULL", "NOT NULL", "NOT NULL", "NOT NULL", "NOT NULL", "-"],
            "Description": [
                "MongoDB auto-generated unique identifier",
                "User email address",
                "One-time verification code",
                "Verification code expiration time",
                "Whether verified",
                "Creation time",
                "Verification completion time"
            ]
        }
    }
    
    return schemas

def format_worksheet(ws, title):
    """Format worksheet with proper styling - black and white theme"""
    
    # Calculate column widths to total 57.67
    # Distribute proportionally: Field Name (20%), Data Type (20%), Constraints (20%), Description (40%)
    total_width = 57.67
    ws.column_dimensions['A'].width = total_width * 0.20  # Field Name
    ws.column_dimensions['B'].width = total_width * 0.20  # Data Type
    ws.column_dimensions['C'].width = total_width * 0.20  # Constraints
    ws.column_dimensions['D'].width = total_width * 0.40  # Description
    
    # Define fonts - Times New Roman, size 12 as requested
    title_font = Font(name='Times New Roman', size=14, bold=True, color="000000")
    header_font = Font(name='Times New Roman', size=12, bold=True, color="FFFFFF")
    cell_font = Font(name='Times New Roman', size=12, color="000000")
    
    # Define fills and borders - black and white theme
    header_fill = PatternFill(start_color="000000", end_color="000000", fill_type="solid")
    border = Border(
        left=Side(style='thin', color="000000"),
        right=Side(style='thin', color="000000"),
        top=Side(style='thin', color="000000"),
        bottom=Side(style='thin', color="000000")
    )
    
    # Set title
    ws['A1'] = title
    ws['A1'].font = title_font
    ws.merge_cells('A1:D1')
    ws['A1'].alignment = Alignment(horizontal='center')
    
    # Set headers
    headers = ['Field Name', 'Data Type', 'Constraints', 'Description']
    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=3, column=col, value=header)
        cell.font = header_font
        cell.fill = header_fill
        cell.border = border
        cell.alignment = Alignment(horizontal='center', vertical='center')
    
    # Set row height for header
    ws.row_dimensions[3].height = 30

def add_data_to_worksheet(ws, schema_data, start_row=4):
    """Add data to worksheet with proper formatting"""
    
    cell_font = Font(name='Times New Roman', size=12, color="000000")
    border = Border(
        left=Side(style='thin', color="000000"),
        right=Side(style='thin', color="000000"),
        top=Side(style='thin', color="000000"),
        bottom=Side(style='thin', color="000000")
    )
    
    # Get the maximum length of any field list
    max_fields = max(len(schema_data[field]) for field in schema_data.keys())
    
    for row in range(max_fields):
        for col, field_name in enumerate(['Field Name', 'Data Type', 'Constraints', 'Description'], 1):
            cell = ws.cell(row=start_row + row, column=col)
            
            if row < len(schema_data[field_name]):
                cell.value = schema_data[field_name][row]
            
            cell.font = cell_font
            cell.border = border
            cell.alignment = Alignment(vertical='center')
            
            # Wrap text for description column
            if col == 4:
                cell.alignment = Alignment(vertical='center', wrap_text=True)
    
    return start_row + max_fields

def create_excel_file():
    """Create the complete Excel file with all schemas"""
    
    # Create workbook
    wb = Workbook()
    
    # Remove default sheet
    wb.remove(wb.active)
    
    # Get schema data
    schemas = create_schema_data()
    
    # Create worksheets for each collection
    for collection_name, schema_data in schemas.items():
        # Excel sheet name limit is 31 characters
        sheet_name = collection_name[:31]
        ws = wb.create_sheet(title=sheet_name)
        
        # Format worksheet
        format_worksheet(ws, collection_name)
        
        # Add data
        add_data_to_worksheet(ws, schema_data)
    
    # Create summary sheet
    summary_ws = wb.create_sheet(title="Summary", index=0)
    format_worksheet(summary_ws, "PreTech-NIDS Database Collections Summary")
    
    # Add summary data
    summary_data = {
        "Collection Name": list(schemas.keys()),
        "Field Count": [len(schemas[name]["Field Name"]) for name in schemas.keys()],
        "Primary Key": ["_id (ObjectId)" for _ in schemas.keys()],
        "Description": [
            "User management and authentication",
            "ML model detection results and network metadata",
            "Security alerts and threat notifications",
            "Configurable alert generation rules",
            "Alert state change history",
            "Geographic attack source mapping",
            "IP geolocation data cache",
            "PCAP file analysis results",
            "PCAP analysis reports",
            "Password reset functionality",
            "User registration verification"
        ]
    }
    
    # Add summary headers
    summary_ws['A3'] = "Collection Name"
    summary_ws['B3'] = "Field Count"
    summary_ws['C3'] = "Primary Key"
    summary_ws['D3'] = "Description"
    
    # Format summary headers
    header_font = Font(name='Times New Roman', size=12, bold=True, color="FFFFFF")
    header_fill = PatternFill(start_color="000000", end_color="000000", fill_type="solid")
    border = Border(
        left=Side(style='thin', color="000000"),
        right=Side(style='thin', color="000000"),
        top=Side(style='thin', color="000000"),
        bottom=Side(style='thin', color="000000")
    )
    
    for col in range(1, 5):
        cell = summary_ws.cell(row=3, column=col)
        cell.font = header_font
        cell.fill = header_fill
        cell.border = border
        cell.alignment = Alignment(horizontal='center', vertical='center')
    
    # Add summary data
    max_rows = len(summary_data["Collection Name"])
    for row in range(max_rows):
        for col, field in enumerate(['Collection Name', 'Field Count', 'Primary Key', 'Description'], 1):
            cell = summary_ws.cell(row=row + 4, column=col, value=summary_data[field][row])
            cell.font = Font(name='Times New Roman', size=12, color="000000")
            cell.border = border
            cell.alignment = Alignment(vertical='center')
            
            # Wrap text for description column
            if col == 4:
                cell.alignment = Alignment(vertical='center', wrap_text=True)
    
    # Save the file
    output_file = "PreTech-NIDS_Database_Schema.xlsx"
    wb.save(output_file)
    print(f"✅ Excel file created successfully: {output_file}")
    print(f"📊 Total collections: {len(schemas)}")
    print(f"📁 File location: {os.path.abspath(output_file)}")
    print(f"🎨 Formatting applied:")
    print(f"   - Total column width: 57.67 (distributed)")
    print(f"   - Font: Times New Roman, size 12")
    print(f"   - Theme: Black and white")
    print(f"   - Borders: All cells")
    
    return output_file

def create_word_document():
    """Create Word document with all database schemas"""
    
    if not WORD_AVAILABLE:
        print("❌ Word document creation skipped - python-docx not available")
        return None
    
    # Create document
    doc = Document()
    
    # Add title
    title = doc.add_heading('PreTech-NIDS Database Schema Documentation', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add overview
    doc.add_heading('Database Overview', level=1)
    doc.add_paragraph('The PreTech-NIDS system employs MongoDB as the primary database management system, a document-based NoSQL database particularly suitable for handling large volumes of unstructured data and real-time data streams generated by network intrusion detection systems.')
    
    # Get schema data
    schemas = create_schema_data()
    
    # Add each collection with numbered headings
    for idx, (collection_name, schema_data) in enumerate(schemas.items(), 1):
        # Create numbered heading (4.3.3.x)
        heading_text = f"4.3.3.{idx} {collection_name}"
        heading = doc.add_heading(heading_text, level=4)
        
        # Apply custom formatting to heading
        for run in heading.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.color.rgb = None  # Black color
            run.font.italic = False  # No italic
        
        # Add table caption
        caption = doc.add_paragraph(f"Table {29 + idx - 1}: {collection_name} Database Schema Table")
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in caption.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)
            run.font.italic = True
        
        # Create table with proper column widths
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Set table width to 6 inches (approximately 57.67 column width equivalent)
        # Calculate column widths: Description 40%, others based on content
        total_width = 6.0  # inches
        desc_width = total_width * 0.40  # 40% for Description
        other_width = (total_width - desc_width) / 3  # Remaining 60% divided by 3
        
        # Header row
        header_cells = table.rows[0].cells
        headers = ['Field Name', 'Data Type', 'Constraints', 'Description']
        
        for i, header in enumerate(headers):
            header_cells[i].text = header
            header_cells[i].paragraphs[0].runs[0].font.bold = True
            header_cells[i].paragraphs[0].runs[0].font.name = 'Times New Roman'
            header_cells[i].paragraphs[0].runs[0].font.size = Pt(12)
            header_cells[i].paragraphs[0].runs[0].font.color.rgb = None  # Black color
            
            # Set column widths
            if i == 3:  # Description column
                header_cells[i].width = Inches(desc_width)
            else:  # Other columns
                header_cells[i].width = Inches(other_width)
        
        # Add data rows
        max_fields = max(len(schema_data[field]) for field in schema_data.keys())
        for row in range(max_fields):
            row_cells = table.add_row().cells
            for col, field_name in enumerate(['Field Name', 'Data Type', 'Constraints', 'Description']):
                if row < len(schema_data[field_name]):
                    row_cells[col].text = str(schema_data[field_name][row])
                row_cells[col].paragraphs[0].runs[0].font.name = 'Times New Roman'
                row_cells[col].paragraphs[0].runs[0].font.size = Pt(12)
                row_cells[col].paragraphs[0].runs[0].font.color.rgb = None  # Black color
                
                # Set column widths for data rows
                if col == 3:  # Description column
                    row_cells[col].width = Inches(desc_width)
                else:  # Other columns
                    row_cells[col].width = Inches(other_width)
        
        # Add spacing after table
        doc.add_paragraph()
    
    # Save document
    output_file = "PreTech-NIDS_Database_Schema.docx"
    doc.save(output_file)
    print(f"✅ Word document created successfully: {output_file}")
    print(f"📁 File location: {os.path.abspath(output_file)}")
    print(f"🎨 Word formatting applied:")
    print(f"   - Total table width: 6 inches (57.67 equivalent)")
    print(f"   - Description column: 40% width")
    print(f"   - Other columns: auto-sized based on content")
    print(f"   - Headings: 4.3.3.x format with Times New Roman 12pt bold")
    print(f"   - Table captions: Table X format")
    
    return output_file

def create_csv_files():
    """Create CSV files as fallback"""
    
    if not CSV_AVAILABLE:
        print("❌ CSV module not available")
        return None
    
    schemas = create_schema_data()
    
    # Create main schema file
    main_file = "../docs/database_schema/PreTech-NIDS_Database_Schema.csv"
    
    with open(main_file, 'w', newline='', encoding='utf-8') as csvfile:
        writer = csv.writer(csvfile)
        
        # Write header
        writer.writerow(['Collection', 'Field Name', 'Data Type', 'Constraints', 'Description'])
        
        # Write data for each collection
        for collection_name, schema_data in schemas.items():
            max_fields = max(len(schema_data[field]) for field in schema_data.keys())
            
            for i in range(max_fields):
                row = [
                    collection_name if i == 0 else "",  # Only show collection name in first row
                    schema_data["Field Name"][i] if i < len(schema_data["Field Name"]) else "",
                    schema_data["Data Type"][i] if i < len(schema_data["Data Type"]) else "",
                    schema_data["Constraints"][i] if i < len(schema_data["Constraints"]) else "",
                    schema_data["Description"][i] if i < len(schema_data["Description"]) else ""
                ]
                writer.writerow(row)
            
            # Add empty row between collections
            writer.writerow([])
    
    print(f"✅ CSV file created: {main_file}")
    print("💡 Open in Excel and format manually")

if __name__ == "__main__":
    try:
        if OPENPYXL_AVAILABLE:
            excel_file = create_excel_file()
            if WORD_AVAILABLE:
                word_file = create_word_document()
                print(f"\n🎉 All files created successfully!")
                print(f"📊 Excel: {excel_file}")
                print(f"📝 Word: {word_file}")
            else:
                print(f"\n🎉 Excel file created successfully!")
                print(f"📊 Excel: {excel_file}")
        else:
            create_csv_files()
    except Exception as e:
        print(f"❌ Error: {e}")
        print("🔄 Falling back to CSV format...")
        create_csv_files()
        import traceback
        traceback.print_exc()
